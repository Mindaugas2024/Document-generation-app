import sys
from PyQt5.QtWidgets import (QApplication, QWidget, QPushButton, QProgressBar, QMainWindow, QAction, QComboBox,
                             QTableWidget, QVBoxLayout, QInputDialog, QFileDialog, QTableWidgetItem, QTextEdit,
                             QHBoxLayout, QCheckBox, QMessageBox)
import shutil
import os
from docx import Document
from database import (create_table_saskaitu_duomenys, create_table_sablonai, insert_data_sablonai, atvaizduoti_sablonus,
                      read_excel, atvaizduoti_saskaitas, insert_data_saskaitos, drop_saskaitos_table)
from docx2pdf import convert

# prisijungimai i duomenu baze
db_params = {
    "host": "localhost",
    "database": "automationApp",
    "user": "postgres",
    "password": "123",
    "port": "5432"
}

class Application(QMainWindow):
    def __init__(self):
        super().__init__()
        self.sablonu_list = []
        self.saskaitu_list = []
        self.sablonu_sarasas = None
        self.saskaitu_sarasas = None
        # self.fileEdit = None
        self.initUI()

    # aplikacijos pagrindinis meniu
    def initUI(self):
        self.showMaximized()
        self.setWindowTitle('Dokumentų Generavimo Sistema')

        # virsutine juosta
        menu_juosta = self.menuBar()

        # pasirinkimai virsutineje juostoje ir susijusiu operaciju mygtukai
        sablonai_menu = menu_juosta.addMenu('Šablonai')
        sablonu_sarasas_button = QAction('Šablonu Sąrašas', self)
        ikelti_sablona_button = QAction('Ikelti Nauja šabloną', self)
        duomenys_menu = menu_juosta.addMenu('Duomenys')
        saskaitu_sarasas_button = QAction('Saskaitų sąrašas', self)
        sutarciu_sarasas_button = QAction('Sutarčių sąrašas', self)

        sablonu_sarasas_button.triggered.connect(self.sablonuSarasas)
        ikelti_sablona_button.triggered.connect(self.ikeltiSablona)
        saskaitu_sarasas_button.triggered.connect(self.saskaituSarasas)

        sablonai_menu.addAction(sablonu_sarasas_button)
        sablonai_menu.addAction(ikelti_sablona_button)
        duomenys_menu.addAction(saskaitu_sarasas_button)
        duomenys_menu.addAction(sutarciu_sarasas_button)

    def sablonuSarasas(self):
        try:

            # jeigu ekrane yra rodomas saskaitu sarasas ji uzdarysime, tam kad rodyti sablonu sarasa
            # tai pat sukuriame jeigu dar neturime duomenu bazeje table kuris laikys musu sablonu duomenis
            # ir pasiemame informacija is db
            create_table_sablonai(db_params)
            self.sablonu_list = atvaizduoti_sablonus(db_params)
            #print(self.sablonu_list)

            # suformuojame lentele sarasui atvaizduoti
            self.sablonu_sarasas = QTableWidget(self)
            self.sablonu_sarasas.setColumnCount(3)
            self.sablonu_sarasas.setHorizontalHeaderLabels(['Pavadinimas', 'Tipas', ''])

            # priskiriame lentelei tiek eiluciu kiek sablonu sarase yra sablonu
            self.sablonu_sarasas.setRowCount(len(self.sablonu_list))

            # pagal kiekviena eilute ikeliame duomenis i atitinkamus stulpelius
            # sablonas: saraso eilute, stulpelis: kokia butent informacija is saraso eilutes bus ideta lenteles
            # stulpeliuose
            for sablonas, stulpelis in enumerate(self.sablonu_list):
                sablono_name = stulpelis["pavadinimas"]
                sablono_tipas = stulpelis["tipas"]

                # i lentele idedame paimtus saraso duomenis
                self.sablonu_sarasas.setItem(sablonas, 0, QTableWidgetItem(sablono_name))
                self.sablonu_sarasas.setItem(sablonas, 1, QTableWidgetItem(sablono_tipas))

                # kiekvienai eilutei lentele pridedame mygtuka modifikuotis kuris mums leis pasirinkti kuri sablona
                # koreguosime
                modifikuoti_button = QPushButton('Modifikuoti', self)
                modifikuoti_button.clicked.connect(lambda checked,row=sablonas: self.modifikuotiSablona(row))
                # print(sablonas)
                self.sablonu_sarasas.setCellWidget(sablonas, 2, modifikuoti_button)

            sablonai_layout = QVBoxLayout()
            sablonai_layout.addWidget(self.sablonu_sarasas)
            central_widget = QWidget(self)
            self.setCentralWidget(central_widget)
            central_widget.setLayout(sablonai_layout)

        except Exception as e:
            print(e)


    def ikeltiSablona(self):
        global naujasis_failo_pavadinimas
        try:
            data = []
            dok_tipai = ['Saskaita', 'Sutartis']

            # pirmasis popup langelis kuriame reikia pasirinkti ikeliamo sablono tipa
            tipas, ok = QInputDialog.getItem(self, 'Ikelti naują šabloną', 'Pasirinkite dokumento tipą:', dok_tipai, 0,

                                            False)
            # antrasis popup kuris praso ivesti sablono pavadinima
            if ok and tipas:
                pavadinimas, ok = QInputDialog.getText(self, 'Ikelti naują šabloną', 'Įveskite pavadinimą:')

                # treciasis popup yra pasirinkimas sablono failo kuri norime isikelti
                if ok and pavadinimas:
                    failo_kelias, _ = QFileDialog.getOpenFileName(self, 'Pasirinkite Failą', '', 'Visi failai (*)')
                    if failo_kelias:

                        # direktorija kurioje bus laikomas sablonas
                        issaugojimo_direktorija = ('C:/Users/minde/Documents/GitHub/PYTHON/Automatizuota Dokumentų '
                                                   'Generavimo Sistema/Sablonai')
                        failo_pavadinimas = os.path.basename(failo_kelias)

                        # perkeliame su shutil musu pasirinkta faila i sablonu direktorija
                        shutil.copy(failo_kelias, issaugojimo_direktorija)

                        # pervadiname faila pagal ivestus duomenis popupe
                        keiciamas_failo_pavadinimas = (issaugojimo_direktorija + '/' + failo_pavadinimas)
                        if "pdf" in failo_pavadinimas:
                            naujasis_failo_pavadinimas = (issaugojimo_direktorija + '/' + pavadinimas + ".pdf")
                        if "docx" in failo_pavadinimas:
                            naujasis_failo_pavadinimas = (issaugojimo_direktorija + '/' + pavadinimas + ".docx")
                        os.rename(keiciamas_failo_pavadinimas, naujasis_failo_pavadinimas)
                        # print(naujasis_failo_pavadinimas)

                        # visus naujus duomenis ikeliame i duomenu baze
                        data.append({"pavadinimas": pavadinimas, "tipas": tipas, "failo_kelias":
                            naujasis_failo_pavadinimas})
                        #print(data)
                        insert_data_sablonai(data,db_params)
            self.sablonuSarasas()
        except Exception as e:
            print(e)

    def modifikuotiSablona(self, row_index):
        try:
            # neberodome saraso ekrane
            self.fileEdit = QTextEdit(self)

            # gauname indeksa mygtuko kuris buvo paspaustas ir su indeksu pasiemame sablono kelia kuri redaguosime

            print(row_index)
            self.modifikuojamas_failas = self.sablonu_list[row_index]["failo_kelias"]
            print(self.modifikuojamas_failas)

            # pagal kelia atsidarome faila ir su for pasiemame ir atvaizduojame teksta kuri galesime koreguoti
            dokumentas = Document(self.modifikuojamas_failas)
            tekstas = ""
            teksto_sarasas = []
            for paragraph in dokumentas.paragraphs:
                line = paragraph.text + '\n'
                tekstas += line
                teksto_sarasas.append(line)

            # for table in dokumentas.tables:
            #     naujas_listas = []
            #     for table_row in table.rows:
            #
            #         for cell in table_row.cells:
            #             cell_tekstas = cell.text.strip() + '\t'
            #             if cell_tekstas not in naujas_listas:
            #
            #                 tekstas += cell_tekstas + '\t'
            #                 langelis = cell_tekstas + '\t'
            #                 naujas_listas.append(cell)
            #                 teksto_sarasas.append(langelis)
            #         # window = [cell.text for cell in row.cells]
            #             # window = cell.text + '\t'
            #             # print(window)
            #             # tekstas += window
            #             # teksto_sarasas.append(window)
            #         # window = '\n'
            #         # print(window)
            #         tekstas += '\n'
            #         # teksto_sarasas.append(window)
            #     # window = '\n'
            #     # tekstas += window
            #     # teksto_sarasas.append(window)
            #
            # print(teksto_sarasas)
            # self.fileEdit.setPlainText(tekstas)

            for table in dokumentas.tables:
                previous_rows_texts = []
                for table_row in table.rows:
                    row_text = ''.join(cell.text.strip() for cell in table_row.cells)
                    if row_text and row_text not in previous_rows_texts:
                        for cell in table_row.cells:
                            tekstas += cell.text.strip() + '\t'
                        tekstas += '\n'
                        previous_rows_texts.append(row_text)
                    tekstas += '\n'
                tekstas += '\n'

            self.fileEdit.setPlainText(tekstas)


            # buttonu layout
            button_layout = QHBoxLayout()
            issaugoti_button = QPushButton("Issaugoti", self)
            issaugoti_button.clicked.connect(self.atnaujintiSablona)
            atsaukti_button = QPushButton("Atsaukti", self)
            #atsaukti_button.clicked.connect(self.fileEdit.hide())
            button_layout.addWidget(issaugoti_button)
            button_layout.addWidget(atsaukti_button)

            # pagrindinis layout ir buttonu layout pridejimas
            modifikacijos_layout = QVBoxLayout()
            modifikacijos_layout.addWidget(self.fileEdit)
            modifikacijos_layout.addLayout(button_layout)
            central_widget = QWidget(self)
            self.setCentralWidget(central_widget)
            central_widget.setLayout(modifikacijos_layout)
        except Exception as e:
            print(e)


    def atnaujintiSablona(self):
        try:
            print("veikia")
            # atnaujiname faila su naujais irasais, irasa fiksuojame pagal eilute
            dokumentas = Document(self.modifikuojamas_failas)
            for element in dokumentas.element.body:
                dokumentas.element.body.remove(element)
            naujas_teksto_sarasas = []
            for eilute in self.fileEdit.toPlainText().split('\n'):
                dokumentas.add_paragraph(eilute)
                naujas_teksto_sarasas.append(eilute)

            dokumentas.save(self.modifikuojamas_failas)
            print(naujas_teksto_sarasas)
        except Exception as e:
            print(e)

    def saskaituSarasas(self):
        try:
            self.sablonuSarasas()
            # toks pat veikimas kaip ir su sablonu sarasu. Isjungiame jeigu rodomas sablonu sarasas, sukuriam table
            # jeigu neturime ir is db pasiemame duomenis ju atvaizdavimui
            create_table_saskaitu_duomenys(db_params)
            self.saskaitu_list = atvaizduoti_saskaitas(db_params)
            # print(self.saskaitu_list)

            self.saskaitu_sarasas = QTableWidget(self)
            self.saskaitu_sarasas.setColumnCount(20)
            self.saskaitu_sarasas.setHorizontalHeaderLabels(['Pazymeti', 'Data', 'Serija', 'Numeris', 'Pardavejo imone',
                                                             'Pardavejo adresas', 'Pardavejo kodas',
                                                             'Pardavejo PVM kodas', 'Pirkejo imone', 'Pirkejo adresas',
                                                             'Pirkejo kodas', 'Pirkejo PVM kodas', 'Preke', 'Mat vnt',
                                                             'Kiekis', 'Kaina be PVM', 'Suma be PVM', 'PVM proc',
                                                             'PVM suma', 'Suma'])
            self.saskaitu_sarasas.setRowCount(len(self.saskaitu_list))

            # pirmasis for pasiema visus saskaitos duomenis pagal eilute sarase, o su antruoju for is kiekvieno eilutes
            # langelio pasirenkame informacija is eiles i atvaizduojame ta informacija table
            for saskaita, stulpelis in enumerate(self.saskaitu_list):
                check_box = QCheckBox()
                self.saskaitu_sarasas.setCellWidget(saskaita, 0, check_box)
                for langelis, (kintamasis, verte) in enumerate(stulpelis.items()):
                    self.saskaitu_sarasas.setItem(saskaita, (langelis + 1), QTableWidgetItem(verte))


            button_layout = QHBoxLayout()
            pasirinkti_location_button = QPushButton('Ikelti duomenis')
            prideti_eilute_button = QPushButton('Ivesti duomenis')
            atnaujinti_button = QPushButton("Atnaujinti", self)
            # formation_bar = QProgressBar()
            # formation_bar.setFixedWidth(305)
            # formation_bar.setValue(0)
            self.pasirinkti_sablona_button = QComboBox(self)
            self.pasirinkti_sablona_button.addItem("Pasirinkti Šabloną")
            for sablonas in self.sablonu_list:
                if sablonas["tipas"] == "Saskaita":
                    self.pasirinkti_sablona_button.addItem(sablonas["pavadinimas"])

            formuoti_button = QPushButton("Formuoti Dokumentus", self)

            pasirinkti_location_button.clicked.connect(self.dokumento_pasirinkimas)
            prideti_eilute_button.clicked.connect(self.nauja_saskaitos_eilute)
            formuoti_button.clicked.connect(self.dokumentu_formavimas)
            atnaujinti_button.clicked.connect(self.atnaujinti_saskaitu_sarasa)
            button_layout.addWidget(pasirinkti_location_button)
            button_layout.addWidget(prideti_eilute_button)
            button_layout.addWidget(atnaujinti_button)
            button_layout.addWidget(self.pasirinkti_sablona_button)
            button_layout.addWidget(formuoti_button)
            # button_layout.addWidget(formation_bar)

            saskaitu_layout = QVBoxLayout()
            saskaitu_layout.addLayout(button_layout)
            saskaitu_layout.addWidget(self.saskaitu_sarasas)
            central_widget = QWidget(self)
            self.setCentralWidget(central_widget)
            central_widget.setLayout(saskaitu_layout)

        except Exception as e:
            print(e)

    def dokumento_pasirinkimas(self):
        try:
            # pasirenkame dokumenta is kompiuterio (.xlsx) kurio duomenis bus atvaizduojame saskaitu sarase
            failo_kelias, _ = QFileDialog.getOpenFileName(self, 'Pasirinkite Failą', '', 'Visi failai (*)')
            if failo_kelias:
                read_excel(failo_kelias, db_params)
            self.saskaituSarasas()
        except Exception as e:
            print(e)

    def nauja_saskaitos_eilute(self):
        try:
            # sis metodas sukuria apacioje tuscia nauja eilute informacijai ivesti
            check_box = QCheckBox()
            paskutine_eilute = self.saskaitu_sarasas.rowCount()
            self.saskaitu_sarasas.insertRow(paskutine_eilute)
            self.saskaitu_sarasas.setCellWidget((paskutine_eilute), 0, check_box)
        except Exception as e:
            print(e)

    def atnaujinti_saskaitu_sarasa(self):
        try:
            # atnaujinant sarasa, susirenkame visus matomus duomenis i sarasa, dropinam database ir
            # sukuriam nauja ir ikeliame naujus duomenis
            data = []
            for row in range(self.saskaitu_sarasas.rowCount()):
                eilutes_data = []
                for collum in range(1, self.saskaitu_sarasas.columnCount()):
                    langelis = self.saskaitu_sarasas.item(row, collum)
                    eilutes_data.append(langelis.text())
                data.append(eilutes_data)
            drop_saskaitos_table(db_params)
            create_table_saskaitu_duomenys(db_params)
            insert_data_saskaitos(data, db_params)
            # print(data)
            self.saskaituSarasas()
        except Exception as e:
            # jeigu kuris nors langelis tuscias gausime klaida
            QMessageBox.critical(self, 'Klaida', 'Ne visi langeliai uzpildyti.')
            print(e)

    def dokumentu_formavimas(self):
        try:
            # prasome vartotojo pasirinkti formavimo tipa
            visa_pasirinkta_info = []
            formavimo_tipai = ["docx", "pdf"]
            # print(self.pasirinkti_sablona_button.currentText())
            if not self.pasirinkti_sablona_button.currentIndex() == 0:
                tipas, ok = QInputDialog.getItem(self, 'Formatas', 'Pasirinkite dokumento formata:', formavimo_tipai,
                                                 0, False)
                sablono_kelias = self.sablonu_list[self.pasirinkti_sablona_button.currentIndex()]["failo_kelias"]
                # print(sablono_kelias)
                # patikriname ar pazymeta kuri nors saskaita
                checked_boxes = [row for row in range(self.saskaitu_sarasas.rowCount())
                                 if self.saskaitu_sarasas.cellWidget(row, 0).isChecked()]

                # graziname antgal i sarasa jeigu nera pazymetu saskaitu
                if not checked_boxes:
                    QMessageBox.warning(self, 'Klaida', 'Nepasirinktos sąskaitos.')
                    return

                for checked_box in checked_boxes:
                    eilutes_duomenis = [self.saskaitu_sarasas.item(checked_box, col).text()
                                     for col in range(1, self.saskaitu_sarasas.columnCount())]
                    visa_pasirinkta_info.append(eilutes_duomenis)


                self.docx_suformavimas(sablono_kelias, visa_pasirinkta_info, tipas)

            else:
                QMessageBox.warning(self, 'Klaida', 'Nepasirinktas šablonas')
        except Exception as e:
            print(e)

    def docx_suformavimas(self, failo_kelias, visa_pasirinkta_info, tipas):
        try:
            # nurodome kur issaugome failus, ir placeholderius kuriu ieskos
            issaugojimas = ('C:/Users/minde/Documents/GitHub/PYTHON/Automatizuota Dokumentų Generavimo Sistema'
                            '/Saskaitos docx/')
            placeholderis = ["{data}", "{serija}", "{numeris}", "{pardavejo imone}", "{pardavejo adresas}",
                             "{pardavejo kodas}", "{pardavejo pvm kodas}", "{pirkejo imone}", "{pirkejo adresas}",
                             "{pirkejo kodas}", "{pirkejo pvm kodas}", "{preke}", "{mat vnt}", "{kiekis}",
                             "{kaina be pvm}", "{suma be pvm}", "{pvm}", "{pvm suma}", "{suma}"]
            # print(visa_pasirinkta_info)

            for saskaita in visa_pasirinkta_info:
                dokumentas = Document(failo_kelias)
                # print(saskaita)
                created_filename = f"{issaugojimas}Saskaita_{saskaita[2]}.docx"
                # print(created_filename)
                for paragraph in dokumentas.paragraphs:
                    for index, placeholder in enumerate(placeholderis):
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, saskaita[index])

                for table in dokumentas.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for index, placeholder in enumerate(placeholderis):
                                if placeholder in cell.text:
                                    cell.text = cell.text.replace(placeholder, saskaita[index])

                dokumentas.save(created_filename)
                if tipas == "pdf":
                    self.pdf_suformavimas(created_filename)
                    # print("vykdom pdf")

        except Exception as e:
            print(e)

    def pdf_suformavimas(self, directory):
        try:
            # pdf failus suformuojame is docx failu
            input_directory = directory
            output_directory = ('C:/Users/minde/Documents/GitHub/PYTHON/Automatizuota Dokumentų Generavimo Sistema'
                                '/saskaitos PDF/')
            print(output_directory)
            convert(input_directory, output_directory)
            os.remove(input_directory)
        except Exception as e:
            print(e)


def main():
    app = QApplication(sys.argv)
    ex = Application()
    ex.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()